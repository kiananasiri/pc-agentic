import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def set_p_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def format_run_rtl(r, font_name="Vazirmatn", font_size=None, bold=None):
    r.font.name = font_name
    if font_size:
        r.font.size = Pt(font_size)
    if bold is not None:
        r.bold = bold
    
    rPr = r._element.get_or_add_rPr()
    bidi = OxmlElement('w:rtl')
    bidi.set(qn('w:val'), '1')
    rPr.append(bidi)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def update_report_live_data():
    file_path = 'Bachelor_proj_report.docx'
    doc = docx.Document(file_path)

    # 1. Update Abstract to mention Live Data Auto-Fetching
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and '۱. چکیده' in p.text:
            abstract_p = doc.paragraphs[i + 1]
            abstract_p.text = ""
            r_abs = abstract_p.add_run(
                "این گزارش به بررسی معماری، طراحی و پیاده‌سازی پروژه دستیار هوشمند بازار کریپتو (دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی) می‌پردازد. "
                "این پروژه پیش‌تر برای پلتفرم parsiancrpto.com طراحی و توسعه داده شده بود و تمرکز اصلی آن بر روی ایجاد یک معماری مبتنی بر ایجنت (Agentic Architecture) "
                "با بهره‌گیری از فریم‌ورک Agno، پایگاه داده PostgreSQL و Redis، و سیستم دریافت خودکار و زنده داده‌های بازار (Live Market & On-Chain Auto-Fetching) است "
                "تا بتواند به صورت هوشمندانه، لحظه‌ای و پایدار به سوالات عمومی، تحلیل‌های تکنیکال و داده‌های درون زنجیره‌ای (On-Chain) پاسخ دهد. "
                "همچنین در بخش فرانت‌اند، از فریم‌ورک Astro برای پیاده‌سازی یک رابط کاربری مدرن، تاریک و راست‌چین استفاده شده است."
            )
            format_run_rtl(r_abs)
            set_p_rtl(abstract_p)
            break

    # 2. Update/Append Section 3.5 for Live Data Auto-Fetching mechanism
    found_35 = False
    for i, p in enumerate(doc.paragraphs):
        if '۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها' in p.text:
            found_35 = True
            p_text = doc.paragraphs[i + 1]
            p_text.text = ""
            r_35 = p_text.add_run(
                "یکی از کلیدی‌ترین قابلیت‌های سیستم، عدم اتکا به داده‌های قدیمی و توانایی دریافت خودکار و لحظه‌ای داده‌های بازار کریپتو (Live Market Data) می‌باشد. چگونگی عملکرد این مکانیزم به شرح زیر است:\n"
                "۱. شناسایی خودکار حالت تحلیل (Auto-Detection): زمانی که کاربر پرسشی مرتبط با تحلیل تکنیکال (مانند تحلیل کندل‌ها و اندیکاتورهای یک نماد مثل BTCUSDT) یا تحلیل آن‌چین (مانند جریان پول هوشمند و رفتار ماینرها) مطرح می‌کند، سرویس MarketDataService به صورت خودکار فعال می‌گردد.\n"
                "۲. فراخوانی APIهای لحظه‌ای (Live Data Fetching): در صورت عدم ارائه فایل نمونه توسط کاربر، سیستم به طور مستقیم و لحظه‌ای آخرین داده‌های قیمت، تاریخچه کندل‌های تایم‌فریم مربوطه و اندیکاتورهای محاسبه‌شده (شامل RSI, MACD, Aroon, CCI, Moving Averages) و همچنین شاخص‌های آن‌چین را استخراج می‌نماید.\n"
                "۳. تزریق به ابزارهای ایجنت (ToolKit Injection): داده‌های زنده دریافت‌شده به صورت ساختاریافته (JSON) به ابزارهای تحلیلی ایجنت (مانند TAToolKit و FundamentalToolKit) تزریق می‌شوند. ایجنت Agno با پردازش این داده‌های زنده، اقدام به شناسایی سطوح حمایت/مقاومت، روند کلی و صدور سیگنال‌های دقیق معاملاتی (ورود، حد سود TP، حد ضرر SL و نسبت R:R) می‌نماید."
            )
            format_run_rtl(r_35)
            set_p_rtl(p_text)
            break

    if not found_35:
        h_35 = doc.add_heading('۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها (Live Data Auto-Fetching)', level=2)
        h_35.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_35 = doc.add_paragraph()
        p_35.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_35 = p_35.add_run(
            "یکی از کلیدی‌ترین قابلیت‌های سیستم، عدم اتکا به داده‌های قدیمی و توانایی دریافت خودکار و لحظه‌ای داده‌های بازار کریپتو (Live Market Data) می‌باشد. چگونگی عملکرد این مکانیزم به شرح زیر است:\n"
            "۱. شناسایی خودکار حالت تحلیل (Auto-Detection): زمانی که کاربر پرسشی مرتبط با تحلیل تکنیکال (مانند تحلیل کندل‌ها و اندیکاتورهای یک نماد مثل BTCUSDT) یا تحلیل آن‌چین (مانند جریان پول هوشمند و رفتار ماینرها) مطرح می‌کند، سرویس MarketDataService به صورت خودکار فعال می‌گردد.\n"
            "۲. فراخوانی APIهای لحظه‌ای (Live Data Fetching): در صورت عدم ارائه فایل نمونه توسط کاربر، سیستم به طور مستقیم و لحظه‌ای آخرین داده‌های قیمت، تاریخچه کندل‌های تایم‌فریم مربوطه و اندیکاتورهای محاسبه‌شده (شامل RSI, MACD, Aroon, CCI, Moving Averages) و همچنین شاخص‌های آن‌چین را استخراج می‌نماید.\n"
            "۳. تزریق به ابزارهای ایجنت (ToolKit Injection): داده‌های زنده دریافت‌شده به صورت ساختاریافته (JSON) به ابزارهای تحلیلی ایجنت (مانند TAToolKit و FundamentalToolKit) تزریق می‌شوند. ایجنت Agno با پردازش این داده‌های زنده، اقدام به شناسایی سطوح حمایت/مقاومت، روند کلی و صدور سیگنال‌های دقیق معاملاتی (ورود، حد سود TP، حد ضرر SL و نسبت R:R) می‌نماید."
        )
        format_run_rtl(r_35)
        set_p_rtl(p_35)

    doc.save('Bachelor_proj_report.docx')
    print("Bachelor_proj_report.docx successfully updated with live data details.")

if __name__ == '__main__':
    update_report_live_data()
