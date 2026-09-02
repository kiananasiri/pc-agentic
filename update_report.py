import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def set_p_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def format_run_rtl(r, font_name="Vazirmatn", font_size=None, bold=None):
    r.font.name = font_name
    if font_size:
        r.font.size = Pt(font_size)
    if bold is not None:
        r.bold = bold
    
    rPr = r._element.get_or_add_rPr()
    bidi = OxmlElement('w:rtl')
    bidi.set(qn('w:val'), '1')
    rPr.append(bidi)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def update_report_database():
    file_path = 'Bachelor_proj_report.docx'
    doc = docx.Document(file_path)

    # 1. Update Abstract to explicitly include PostgreSQL & Redis database
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and '۱. چکیده' in p.text:
            abstract_p = doc.paragraphs[i + 1]
            abstract_p.text = ""
            r_abs = abstract_p.add_run(
                "این گزارش به بررسی معماری، طراحی و پیاده‌سازی پروژه دستیار هوشمند بازار کریپتو (دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی) می‌پردازد. "
                "این پروژه پیش‌تر برای پلتفرم parsiancrpto.com طراحی و توسعه داده شده بود و تمرکز اصلی آن بر روی ایجاد یک معماری مبتنی بر ایجنت (Agentic Architecture) "
                "با بهره‌گیری از فریم‌ورک Agno و پایگاه داده PostgreSQL و Redis است تا بتواند به صورت هوشمندانه و پایدار به سوالات عمومی، تحلیل‌های تکنیکال و داده‌های درون زنجیره‌ای (On-Chain) پاسخ دهد. "
                "همچنین در بخش فرانت‌اند، از فریم‌ورک Astro برای پیاده‌سازی یک رابط کاربری مدرن، تاریک و راست‌چین استفاده شده است."
            )
            format_run_rtl(r_abs)
            set_p_rtl(abstract_p)
            break

    # 2. Update System Architecture Overview to explicitly include database layer
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and '۲. معماری کلی سیستم' in p.text:
            overview_p = doc.paragraphs[i + 1]
            overview_p.text = ""
            r_ov = overview_p.add_run(
                "سیستم از دو بخش اصلی تشکیل شده است: یک بک‌اند قدرتمند مبتنی بر پایتون، جنگو و فریم‌ورک Agno متصل به پایگاه داده PostgreSQL و Redis، و یک فرانت‌اند سریع و مدرن بر پایه Astro. "
                "این سیستم در ابتدا جهت ارائه خدمات تحلیلی هوشمند به کاربران سامانه parsiancrpto.com طراحی و آماده‌سازی گردید. "
                "ارتباط بین این دو بخش از طریق APIهای RESTful (به طور خاص مسیر /api/chat) صورت می‌گیرد. درخواست‌های کاربر دریافت شده، داده‌های متصل شده (در صورت وجود) پردازش شده و پس از ذخیره‌سازی و مدیریت در پایگاه داده، به ایجنت‌های مربوطه جهت تولید پاسخ تحویل داده می‌شوند."
            )
            format_run_rtl(r_ov)
            set_p_rtl(overview_p)
            break

    # 3. Enhance Section 3.4 (Database section) with full PostgreSQL schema details
    for i, p in enumerate(doc.paragraphs):
        if '۳.۴. پایگاه داده' in p.text:
            db_p = doc.paragraphs[i + 1]
            db_p.text = ""
            r_db = db_p.add_run(
                "جهت مدیریت نشست‌های کاری (Conversations)، ذخیره‌سازی پیام‌ها، نگه‌داری متاداده‌های تحلیلی و ثبت بازخوردهای کاربران، بک‌اند سیستم از پایگاه داده رابطه‌ای PostgreSQL به همراه ORM قدرتمند جنگو استفاده می‌کند. مهم‌ترین مدل‌های ذخیره‌سازی طراحی‌شده عبارتند از:\n"
                "- Conversation: نگه‌داری شناسه یکتای نشست (UUID)، شناسه کاربر (user_id)، وضعیت آرشیو (is_archive) و زمان‌بندی ایجاد و به‌روزرسانی.\n"
                "- Message: مدل اصلی ذخیره‌سازی پیام‌ها شامل متن پرسش کاربر (question)، پاسخ تولیدشده توسط ایجنت (answer)، قیمت محاسباتی به دلار (price)، مدل هوش مصنوعی (model مانند gpt-4o, o3-mini)، وضعیت اجرای درخواست (status)، سطح استدلال (reasoning)، مدت زمان تولید پاسخ (answer_generation_duration) و فیلدهای ساختاریافته JSON (تحت عنوان analysis برای تحلیل تکنیکال و analysis_fundamental برای تحلیل آن‌چین و جریان پول هوشمند).\n"
                "- Feedback: ثبت بازخوردهای کیفی کاربران (لایک / دیس‌لایک) به همراه شناسه کاربر و کلید خارجی متصل به پیام جهت سنجش دقیق کیفیت ایجنت‌ها.\n"
                "- ProblemReport: سیستم ثبت و پیگیری گزارش خطاهای کاربران به همراه وضعیت بررسی (Open/Close) و پاسخ مدیر سیستم (close_admin_id).\n"
                "همچنین پایگاه داده کلید-مقدار Redis به عنوان Broker پیام‌رسان برای مدیریت وظایف پیش‌زمینه (Asynchronous Tasks) و لایه لایه‌بندی حافظه پنهان (Caching) استفاده گردیده است."
            )
            format_run_rtl(r_db)
            set_p_rtl(db_p)
            break

    doc.save('Bachelor_proj_report.docx')
    print("Bachelor_proj_report.docx database details updated successfully.")

if __name__ == '__main__':
    update_report_database()
