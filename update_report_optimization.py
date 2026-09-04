import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

def set_p_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def format_run_rtl(r, font_name="Vazirmatn", font_size=None, bold=None):
    r.font.name = font_name
    if font_size:
        r.font.size = Pt(font_size)
    if bold is not None:
        r.bold = bold
    
    rPr = r._element.get_or_add_rPr()
    bidi = OxmlElement('w:rtl')
    bidi.set(qn('w:val'), '1')
    rPr.append(bidi)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def update_report():
    file_path = 'Bachelor_proj_report.docx'
    doc = docx.Document(file_path)

    # 1. Update Section 3.3 (Orchestrator)
    for i, p in enumerate(doc.paragraphs):
        if '۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)' in p.text:
            orch_p = doc.paragraphs[i + 1]
            orch_p.text = ""
            r_orch = orch_p.add_run(
                "کامپوننت orchestrator.py وظیفه مسیریابی (Routing) هوشمند درخواست‌ها را بر عهده دارد. تابع run_agno_chat با بررسی دقیق متن پرسش و دیتاسِت‌های ورودی، درخواست کاربر را به ایجنت تخصصی مربوطه (Technical, Fundamental, Code یا Reactive) هدایت می‌کند. "
                "در این لایه، گام‌های روند تفکر ایجنت (Thinking Process) به صورت شفاف استخراج شده و محاسبه دقیق توکن‌های مصرفی و قیمت نهایی پردازش به دلار (مانند $0.000080 برای gpt-4o-mini) بر اساس فرمول‌های قیمت‌گذاری مدل انجام می‌پذیرد."
            )
            format_run_rtl(r_orch)
            set_p_rtl(orch_p)
            break

    # 2. Update Section 3.5 (Live Data Auto-Fetching & Optimization)
    for i, p in enumerate(doc.paragraphs):
        if '۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها' in p.text:
            live_p = doc.paragraphs[i + 1]
            live_p.text = ""
            r_live = live_p.add_run(
                "یکی از کلیدی‌ترین قابلیت‌های سیستم، عدم اتکا به داده‌های قدیمی و توانایی دریافت خودکار و لحظه‌ای داده‌های بازار کریپتو (Live Market Data) می‌باشد. چگونگی عملکرد و بهینه‌سازی این مکانیزم به شرح زیر است:\n"
                "۱. شناسایی هوشمند حالت تحلیل (Query Intent Detection): هنگام طرح پرسش‌های تحلیلی، سرویس MarketDataService کلیدواژه‌های تخصصی تکنیکال (مانند RSI, MACD, کندل، قیمت) و آن‌چین (مانند پول هوشمند، رفتار ماینرها) را بررسی می‌کند.\n"
                "۲. فراخوانی APIهای لحظه‌ای (Live Market Data API): در صورت نیاز به داده، داده‌های لحظه‌ای قیمت و کندل‌های صرافی Binance (پروتوکل API klines) و شاخص‌های آن‌چین CoinGecko استخراج می‌شوند.\n"
                "۳. بهینه‌سازی عدم فراخوانی در گفتگوهای عمومی (General Chat Optimization): در نسخه قبلی، عدم کشف کلیدواژه باعث فراخوانی پیش‌فرض داده‌های بیت‌کوین برای پرسش‌های ساده (مانند 'hi' یا 'سلام') می‌شد. با بهینه‌سازی صورت‌گرفته در تابع auto_enrich_payloads و افزودن تست واحد test_auto_enrich_payloads_greeting، برای پیام‌های عمومی هیچ فراخوانی API خارجی انجام نگرفته و پرسش مستقیماً توسط Reactive Agent با حداکثر سرعت و بدون هزینه ابزار پاسخ داده می‌شود.\n"
                "۴. تزریق به ابزارهای ایجنت (ToolKit Injection): داده‌های زنده به صورت ساختاریافته (JSON) به ابزارهای تحلیلی ایجنت (TAToolKit و FundamentalToolKit) تزریق می‌شوند تا ایجنت Agno سطوح حمایت/مقاومت، روند بازار و سیگنال‌های معاملاتی (EP, TP, SL, R:R) را استخراج کند."
            )
            format_run_rtl(r_live)
            set_p_rtl(live_p)
            break

    # 3. Update Section 5.1 (Backend Unit Tests) to explicitly mention auto-enrichment greeting test
    for i, p in enumerate(doc.paragraphs):
        if '۵.۱. تست‌های واحد و یکپارچه‌سازی بک‌اند' in p.text:
            test_p = doc.paragraphs[i + 1]
            test_p.text = ""
            r_test = test_p.add_run(
                "جهت حصول اطمینان از صحت عملکرد ایجنت‌ها، ابزارهای تحلیلی و عدم وقوع خطا، مجموعه تست‌های واحد بر پایه فریم‌ورک pytest پیاده‌سازی شده است (شامل ۱۶ تست موفق). مهم‌ترین موارد عبارتند از:\n"
                "- test_orchestrator: اعتبارسنجی محاسبه توکن، قیمت‌گذاری دلار، ساخت متن تاریخچه گفتگو و مسیریابی صحیح ارکستریتور.\n"
                "- test_auto_enrich_payloads_greeting: تست تایید عدم فراخوانی APIهای خارجی Binance و CoinGecko هنگام ارسال پیام‌های احوال‌پرسی ساده (مانند 'hi').\n"
                "- test_ta_tools: تست توابعی مانند get_chart_summary, get_recent_candles و get_indicator_data جهت تایید خروجی ساختاریافته JSON.\n"
                "- test_fundamental_tools: تست توابع get_fundamental_summary و get_metric_data جهت سنجش تحلیل‌های آن‌چین."
            )
            format_run_rtl(r_test)
            set_p_rtl(test_p)
            break

    doc.save(file_path)
    print(f"Report '{file_path}' updated successfully with auto-enrichment optimization details.")

if __name__ == '__main__':
    update_report()
