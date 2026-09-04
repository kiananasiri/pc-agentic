import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define color constants
PINK_COLOR = RGBColor(233, 30, 99)      # Vibrant Pink (#E91E63)
PINK_HEX = "E91E63"
HEADING2_COLOR = RGBColor(30, 58, 138)  # Deep Navy Blue (#1E3A8A)
HEADING2_HEX = "1E3A8A"
BODY_COLOR = RGBColor(33, 37, 41)       # Dark Charcoal (#212529)
BODY_HEX = "212529"
FONT_NAME = "Vazirmatn"

def set_paragraph_rtl(p):
    pPr = p._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def format_run_font(run, font_name=FONT_NAME, size_pt=None, bold=None, color_rgb=None, color_hex=None):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb

    rPr = run._element.get_or_add_rPr()
    
    # Enable RTL for run
    bidi = OxmlElement('w:rtl')
    bidi.set(qn('w:val'), '1')
    rPr.append(bidi)
    
    # Set rFonts for all script variants
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

    # Set explicit XML color if hex provided
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(h)
    
    # Styling based on level
    if level == 1:
        # Heading 1 -> PINK
        for run in h.runs:
            format_run_font(run, font_name=FONT_NAME, size_pt=18, bold=True, color_rgb=PINK_COLOR, color_hex=PINK_HEX)
    elif level == 2:
        # Heading 2 -> Navy Blue
        for run in h.runs:
            format_run_font(run, font_name=FONT_NAME, size_pt=14, bold=True, color_rgb=HEADING2_COLOR, color_hex=HEADING2_HEX)
    else:
        for run in h.runs:
            format_run_font(run, font_name=FONT_NAME, size_pt=12, bold=True, color_rgb=BODY_COLOR, color_hex=BODY_HEX)
            
    return h

def add_styled_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(p)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(8)
    
    r = p.add_run(text)
    format_run_font(r, font_name=FONT_NAME, size_pt=11, color_rgb=BODY_COLOR, color_hex=BODY_HEX)
    return p

def create_report():
    doc = docx.Document()

    # Document-wide styles configuration
    styles = doc.styles
    
    # Configure Normal Style
    normal_style = styles['Normal']
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_COLOR

    # Configure Heading 1 Style -> PINK
    h1_style = styles['Heading 1']
    h1_style.font.name = FONT_NAME
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = PINK_COLOR

    # Configure Heading 2 Style
    h2_style = styles['Heading 2']
    h2_style.font.name = FONT_NAME
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = HEADING2_COLOR

    # Configure Title Style
    title_style = styles['Title']
    title_style.font.name = FONT_NAME
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(15, 23, 42)

    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(title_p)
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(24)
    r_title = title_p.add_run('گزارش معماری و پیاده‌سازی پروژه دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی')
    format_run_font(r_title, font_name=FONT_NAME, size_pt=22, bold=True, color_rgb=RGBColor(15, 23, 42))

    doc.add_paragraph('\n\n')

    prof_p = doc.add_paragraph()
    prof_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(prof_p)
    r_prof = prof_p.add_run('درس هوش مصنوعی - استاد تاج بخش')
    format_run_font(r_prof, font_name=FONT_NAME, size_pt=16, bold=True, color_rgb=RGBColor(51, 65, 85))

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(name_p)
    r_name = name_p.add_run('کیانا نصیری')
    format_run_font(r_name, font_name=FONT_NAME, size_pt=14, bold=True, color_rgb=RGBColor(51, 65, 85))

    term_p = doc.add_paragraph()
    term_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rtl(term_p)
    r_term = term_p.add_run('تابستان ۱۴۰۵')
    format_run_font(r_term, font_name=FONT_NAME, size_pt=13, color_rgb=RGBColor(100, 116, 139))

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_styled_heading(doc, 'فهرست مطالب', level=1)
    
    toc_items = [
        "۱. چکیده",
        "۲. معماری کلی سیستم",
        "۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)",
        "   ۳.۱. ایجنت‌های هوشمند (Agno Agents)",
        "   ۳.۲. ابزارهای تحلیلی (ToolKits)",
        "   ۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)",
        "   ۳.۴. پایگاه داده و ذخیره‌سازی داده‌ها (PostgreSQL & Redis)",
        "   ۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها (Live Data Auto-Fetching)",
        "۴. معماری فرانت‌اند",
        "   ۴.۱. پشته فناوری (Tech Stack)",
        "   ۴.۲. کامپوننت‌های مهم و کلاس‌ها",
        "۵. تست و ارزیابی سیستم (Testing & Quality Assurance)",
        "   ۵.۱. تست‌های واحد و یکپارچه‌سازی بک‌اند (Pytest Suite)",
        "   ۵.۲. تست محاسبات تریدینگ و لوریج",
        "   ۵.۳. تست آداپتور فرانت‌اند و API",
        "۶. نتیجه‌گیری"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(p)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        is_sub = item.startswith("   ")
        r = p.add_run(item)
        format_run_font(
            r,
            font_name=FONT_NAME,
            size_pt=10.5 if is_sub else 11.5,
            bold=not is_sub,
            color_rgb=RGBColor(71, 85, 105) if is_sub else RGBColor(15, 23, 42)
        )

    doc.add_page_break()

    # ==================== SECTION 1 ====================
    add_styled_heading(doc, '۱. چکیده', level=1)
    add_styled_paragraph(
        doc,
        "این گزارش به بررسی معماری، طراحی و پیاده‌سازی پروژه دستیار هوشمند بازار کریپتو (دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی) می‌پردازد. "
        "این پروژه پیش‌تر برای پلتفرم parsiancrpto.com طراحی و توسعه داده شده بود و تمرکز اصلی آن بر روی ایجاد یک معماری مبتنی بر ایجنت (Agentic Architecture) "
        "با بهره‌گیری از فریم‌ورک Agno، پایگاه داده PostgreSQL و Redis، و سیستم دریافت خودکار و زنده داده‌های بازار (Live Market & On-Chain Auto-Fetching) است "
        "تا بتواند به صورت هوشمندانه، لحظه‌ای و پایدار به سوالات عمومی، تحلیل‌های تکنیکال و داده‌های درون زنجیره‌ای (On-Chain) پاسخ دهد. "
        "همچنین در بخش فرانت‌اند، از فریم‌ورک Astro برای پیاده‌سازی یک رابط کاربری مدرن، تاریک و راست‌چین استفاده شده است."
    )

    # ==================== SECTION 2 ====================
    add_styled_heading(doc, '۲. معماری کلی سیستم', level=1)
    add_styled_paragraph(
        doc,
        "سیستم از دو بخش اصلی تشکیل شده است: یک بک‌اند قدرتمند مبتنی بر پایتون، جنگو و فریم‌ورک Agno متصل به پایگاه داده PostgreSQL و Redis، و یک فرانت‌اند سریع و مدرن بر پایه Astro. "
        "این سیستم در ابتدا جهت ارائه خدمات تحلیلی هوشمند به کاربران سامانه parsiancrpto.com طراحی و آماده‌سازی گردید. "
        "ارتباط بین این دو بخش از طریق APIهای RESTful (به طور خاص مسیر /api/chat) صورت می‌گیرد. درخواست‌های کاربر دریافت شده، داده‌های متصل شده (در صورت وجود) پردازش شده و پس از ذخیره‌سازی و مدیریت در پایگاه داده، به ایجنت‌های مربوطه جهت تولید پاسخ تحویل داده می‌شوند."
    )

    # ==================== SECTION 3 ====================
    add_styled_heading(doc, '۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)', level=1)
    add_styled_paragraph(
        doc,
        "هسته هوش مصنوعی این پروژه به جای استفاده از روش‌های ساده و یکپارچه، از رویکرد ایجنت‌محور با بهره‌گیری از فریم‌ورک Agno نسخه 2.1.4 استفاده می‌کند. "
        "این روش باعث می‌شود تا وظایف پیچیده تحلیلی شکسته شده و به ایجنت‌های تخصصی واگذار شوند."
    )

    add_styled_heading(doc, '۳.۱. ایجنت‌های هوشمند (Agno Agents)', level=2)
    add_styled_paragraph(
        doc,
        "در این سیستم چندین ایجنت تخصصی تعریف شده است:\n"
        "- Technical Agent: متخصص تحلیل نمودارها، کندل‌ها و اندیکاتورها (مانند RSI, MACD, Aroon). این ایجنت با دریافت داده‌های سری زمانی، وضعیت بازار را تحلیل کرده و سیگنال ارائه می‌دهد.\n"
        "- Fundamental Agent: متخصص بررسی داده‌های آن‌چین (On-Chain) مانند جریان پول هوشمند (Smart Money Flow) و رفتار ماینرها.\n"
        "- Code Agent: ایجنت برنامه‌نویس برای تولید اسکریپت‌های تحلیلی پایتون (مثلاً با Pandas) جهت محاسبه اندیکاتورها.\n"
        "- Reactive Agent (Web): ایجنت عمومی با قابلیت جستجو در وب (از طریق DuckDuckGo) برای پاسخ به سوالات عمومی و اخبار کریپتو."
    )

    add_styled_heading(doc, '۳.۲. ابزارهای تحلیلی (ToolKits)', level=2)
    add_styled_paragraph(
        doc,
        "هر ایجنت برای انجام وظایف خود به مجموعه‌ای از ابزارها دسترسی دارد. این ابزارها در قالب ToolKit توسعه داده شده‌اند:\n"
        "- TAToolKit: شامل توابعی مانند get_chart_summary، get_recent_candles و get_indicator_data که داده‌های خام را برای ایجنت تکنیکال پردازش می‌کنند.\n"
        "- FundamentalToolKit: شامل توابعی مانند get_fundamental_summary و get_metric_data که برای تحلیل جریان‌های ورودی و خروجی صرافی‌ها و رفتار ماینرها استفاده می‌شوند.\n"
        "- WebTools: دسترسی به موتور جستجوی وب برای استخراج اطلاعات بلادرنگ."
    )

    add_styled_heading(doc, '۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)', level=2)
    add_styled_paragraph(
        doc,
        "کامپوننت orchestrator.py وظیفه مسیریابی (Routing) هوشمند درخواست‌ها را بر عهده دارد. تابع run_agno_chat با بررسی دقیق متن پرسش و دیتاسِت‌های ورودی، درخواست کاربر را به ایجنت تخصصی مربوطه (Technical, Fundamental, Code یا Reactive) هدایت می‌کند. "
        "در این لایه، گام‌های روند تفکر ایجنت (Thinking Process) به صورت شفاف استخراج شده و محاسبه دقیق توکن‌های مصرفی و قیمت نهایی پردازش به دلار (مانند $0.000080 برای gpt-4o-mini) بر اساس فرمول‌های قیمت‌گذاری مدل انجام می‌پذیرد."
    )

    add_styled_heading(doc, '۳.۴. پایگاه داده و ذخیره‌سازی داده‌ها (PostgreSQL & Redis)', level=2)
    add_styled_paragraph(
        doc,
        "جهت مدیریت نشست‌های کاری (Conversations)، ذخیره‌سازی پیام‌ها، نگه‌داری متاداده‌های تحلیلی و ثبت بازخوردهای کاربران، بک‌اند سیستم از پایگاه داده رابطه‌ای PostgreSQL به همراه ORM قدرتمند جنگو استفاده می‌کند. مهم‌ترین مدل‌های ذخیره‌سازی طراحی‌شده عبارتند از:\n"
        "- Conversation: نگه‌داری شناسه یکتای نشست (UUID)، شناسه کاربر (user_id)، وضعیت آرشیو (is_archive) و زمان‌بندی ایجاد و به‌روزرسانی.\n"
        "- Message: مدل اصلی ذخیره‌سازی پیام‌ها شامل متن پرسش کاربر (question)، پاسخ تولیدشده توسط ایجنت (answer)، قیمت محاسباتی به دلار (price)، مدل هوش مصنوعی (model مانند gpt-4o, o3-mini)، وضعیت اجرای درخواست (status)، سطح استدلال (reasoning)، مدت زمان تولید پاسخ (answer_generation_duration) و فیلدهای ساختاریافته JSON (تحت عنوان analysis برای تحلیل تکنیکال و analysis_fundamental برای تحلیل آن‌چین و جریان پول هوشمند).\n"
        "- Feedback: ثبت بازخوردهای کیفی کاربران (لایک / دیس‌لایک) به همراه شناسه کاربر و کلید خارجی متصل به پیام جهت سنجش دقیق کیفیت ایجنت‌ها.\n"
        "- ProblemReport: سیستم ثبت و پیگیری گزارش خطاهای کاربران به همراه وضعیت بررسی (Open/Close) و پاسخ مدیر سیستم (close_admin_id).\n"
        "همچنین پایگاه داده کلید-مقدار Redis به عنوان Broker پیام‌رسان برای مدیریت وظایف پیش‌زمینه (Asynchronous Tasks) و لایه لایه‌بندی حافظه پنهان (Caching) استفاده گردیده است."
    )

    add_styled_heading(doc, '۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها (Live Data Auto-Fetching)', level=2)
    add_styled_paragraph(
        doc,
        "یکی از کلیدی‌ترین قابلیت‌های سیستم، عدم اتکا به داده‌های قدیمی و توانایی دریافت خودکار و لحظه‌ای داده‌های بازار کریپتو (Live Market Data) می‌باشد. چگونگی عملکرد و بهینه‌سازی این مکانیزم به شرح زیر است:\n"
        "۱. شناسایی هوشمند حالت تحلیل (Query Intent Detection): هنگام طرح پرسش‌های تحلیلی، سرویس MarketDataService کلیدواژه‌های تخصصی تکنیکال (مانند RSI, MACD, کندل، قیمت) و آن‌چین (مانند پول هوشمند، رفتار ماینرها) را بررسی می‌کند.\n"
        "۲. فراخوانی APIهای لحظه‌ای (Live Market Data API): در صورت نیاز به داده، داده‌های لحظه‌ای قیمت و کندل‌های صرافی Binance (پروتوکل API klines) و شاخص‌های آن‌چین CoinGecko استخراج می‌شوند.\n"
        "۳. بهینه‌سازی عدم فراخوانی در گفتگوهای عمومی (General Chat Optimization): در نسخه قبلی، عدم کشف کلیدواژه باعث فراخوانی پیش‌فرض داده‌های بیت‌کوین برای پرسش‌های ساده (مانند 'hi' یا 'سلام') می‌شد. با بهینه‌سازی صورت‌گرفته در تابع auto_enrich_payloads و افزودن تست واحد test_auto_enrich_payloads_greeting، برای پیام‌های عمومی هیچ فراخوانی API خارجی انجام نگرفته و پرسش مستقیماً توسط Reactive Agent با حداکثر سرعت و بدون هزینه ابزار پاسخ داده می‌شود.\n"
        "۴. تزریق به ابزارهای ایجنت (ToolKit Injection): داده‌های زنده به صورت ساختاریافته (JSON) به ابزارهای تحلیلی ایجنت (TAToolKit و FundamentalToolKit) تزریق می‌شوند تا ایجنت Agno سطوح حمایت/مقاومت، روند بازار و سیگنال‌های معاملاتی (EP, TP, SL, R:R) را استخراج کند."
    )

    # ==================== SECTION 4 ====================
    add_styled_heading(doc, '۴. معماری فرانت‌اند', level=1)

    add_styled_heading(doc, '۴.۱. پشته فناوری (Tech Stack)', level=2)
    add_styled_paragraph(
        doc,
        "فرانت‌اند این سیستم به طور کامل با فریم‌ورک Astro و با استفاده از آداپتور @astrojs/node برای حالت Server-Side Rendering (SSR) پیاده‌سازی شده است. "
        "این انتخاب به دلیل سرعت بالای بارگذاری و امکان اجرای کدهای سمت سرور (جهت برقراری ارتباط امن با بک‌اند) صورت گرفته است. استایل‌دهی نیز با استفاده از "
        "CSS خام مدرن با تم تاریک (Dark Mode) و کلاس‌های شیشه‌ای (Glassmorphism) انجام شده است."
    )

    add_styled_heading(doc, '۴.۲. کامپوننت‌های مهم و کلاس‌ها', level=2)
    add_styled_paragraph(
        doc,
        "مهم‌ترین بخش‌های پیاده‌سازی شده در فرانت‌اند عبارتند از:\n"
        "- ChatInterface.astro: اصلی‌ترین کامپوننت که شامل فرم ورودی، انتخابگر مدل (Model Selector)، تب‌های انتخاب حالت (عمومی، تکنیکال، بنیادین، کدنویسی) و نمایش پیام‌ها است.\n"
        "- ChatMessage.astro: کامپوننت نمایش حباب‌های چت کاربر و ایجنت. این کامپوننت کلاس‌های راست‌چین (rtl-text) را برای متون فارسی اعمال کرده و قیمت هر درخواست را نیز نشان می‌دهد.\n"
        "- Sidebar.astro: سایدبار پروژه که شامل دکمه ایجاد چت جدید و کارت‌های سناریوهای آماده (تحلیل BTC 1m، جریان پول هوشمند، رفتار ماینرها) با استایل‌های کنتراست بالا (High Contrast) می‌باشد.\n"
        "- api/chat.ts: اندپوینت سمت سرور در Astro که وظیفه برقراری ارتباط (Proxy) با بک‌اند هوش مصنوعی (Agno Orchestrator) را بر عهده دارد."
    )

    # ==================== SECTION 5 ====================
    add_styled_heading(doc, '۵. تست و ارزیابی سیستم (Testing & Quality Assurance)', level=1)

    add_styled_heading(doc, '۵.۱. تست‌های واحد و یکپارچه‌سازی بک‌اند (Pytest Suite)', level=2)
    add_styled_paragraph(
        doc,
        "جهت حصول اطمینان از صحت عملکرد ایجنت‌ها، ابزارهای تحلیلی و کامپوننت هماهنگ‌کننده، یک مجموعه جامع از تست‌های واحد و یکپارچه‌سازی با فریم‌ورک Pytest پیاده‌سازی شده است (شامل ۱۶ تست با موفقیت ۱۰۰٪). مهم‌ترین بخش‌های تست شده عبارتند از:\n"
        "- test_chat_bot.py: ارزیابی تابع اصلی ورودی چت، مدیریت ورودی‌های خالی و شبیه‌سازی (Mocking) پاسخ‌های ایجنت Agno شامل محاسبه توکن‌ها و هزینه درخواست.\n"
        "- test_fundamental_tools.py: تست ابزارهای آن‌چین در FundamentalToolKit مانند خلاصه متریال‌های بنیادی، متدهای لیست‌گیری و استخراج داده‌های جریان پول هوشمند (Smart Money Flow) و خالص جریان صرافی‌ها (Exchange Netflow).\n"
        "- test_orchestrator.py: تست توابع هماهنگ‌کننده شامل تخمین توکن‌ها (estimate_tokens)، قالب‌بندی بافت تاریخچه مکالمه (format_history_context)، محاسبه قیمت دقیق مدل‌های مختلف (gpt-4o-mini, o3-mini, gpt-4.5-preview) و اجرای حالت جایگزین (Fallback).\n"
        "- test_ta_tools.py: سنجش ابزارهای تحلیل تکنیکال TAToolKit شامل خلاصه نمودار، دریافت کندل‌های OHLCV، محاسبه آمار قیمت (Highest High و Lowest Low) و پردازش خروجی اندیکاتورهای RSI, Aroon و MACD.\n"
        "- test_auto_enrich_payloads_greeting: تست تایید عدم فراخوانی APIهای خارجی Binance و CoinGecko هنگام ارسال پیام‌های احوال‌پرسی ساده (مانند 'hi')."
    )

    add_styled_heading(doc, '۵.۲. تست محاسبات تریدینگ و لوریج', level=2)
    add_styled_paragraph(
        doc,
        "یکی از حساس‌ترین بخش‌های پروژه، اعتبارسنجی محاسبات تریدینگ است. در این راستا تست‌های دقیق برای توابع ریاضی زیر اجرا گردید:\n"
        "- محاسبه ضریب لوریج بر اساس درصد حد زیان: اعتبارسنجی فرمول LV = 90 / %SL و گرد کردن آن به سمت پایین.\n"
        "- محاسبه نسبت ریسک به ریوارد (Risk-Reward Ratio): ارزیابی نسبت سود به زیان برای پوزیشن‌های خرید (Long) و فروش (Short) بر اساس قیمت ورود، حد زیان و اهداف سود (TP1..N).\n"
        "- اعتبارسنجی پارامترهای معاملاتی (validate_trading_calculations): کنترل حد زیان متناسب با تایم‌فریم (مانند حداقل ۲٪ برای تایم‌فریم ۱۵ دقیقه) و صدور هشدار در صورت ریسک نامناسب."
    )

    add_styled_heading(doc, '۵.۳. تست آداپتور فرانت‌اند و API', level=2)
    add_styled_paragraph(
        doc,
        "در بخش فرانت‌اند Astro، اندپوینت api/chat.ts از نظر سرعت پاسخ‌دهی سمت سرور (SSR)، لایه‌بندی حباب‌های پیام، نمایش کارت‌های سناریو پیش‌فرض و مدیریت خطاهای شبکه مورد ارزیابی قرار گرفت تا تجربه کاربری روان و بلادرنگ برای کاربران تضمین گردد."
    )

    # ==================== SECTION 6 ====================
    add_styled_heading(doc, '۶. نتیجه‌گیری', level=1)
    add_styled_paragraph(
        doc,
        "با مهاجرت از یک سیستم یکپارچه ساده به معماری ایجنت‌محور بر پایه Agno در بک‌اند، بهره‌گیری از پایگاه داده PostgreSQL و Redis جهت مدیریت مطمئن نشست‌ها و کارهای پیش‌زمینه، و استفاده از رابط کاربری سریع مبتنی بر Astro در فرانت‌اند، اکنون سیستم توانایی پردازش هوشمندانه‌تر درخواست‌ها، انتخاب ابزارهای مناسب برای هر نوع پرسش و ارائه تجربه کاربری بومی (فارسی و دارک مود) با کنتراست بالا را دارا می‌باشد."
    )

    doc.save('Bachelor_proj_report.docx')
    print("Report formatted successfully as Bachelor_proj_report.docx with Vazirmatn font and Pink Heading 1!")

if __name__ == '__main__':
    create_report()
