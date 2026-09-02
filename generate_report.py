import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

doc = docx.Document()

# Add Title Page
title = doc.add_heading('گزارش معماری و پیاده‌سازی پروژه PC_BOT_django', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n\n\n')

prof_p = doc.add_paragraph('درس هوش مصنوعی - استاد تاج بخش')
prof_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
prof_p.runs[0].font.size = Pt(16)
prof_p.runs[0].font.bold = True

name_p = doc.add_paragraph('کیانا نصیری')
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.runs[0].font.size = Pt(14)
name_p.runs[0].font.bold = True

term_p = doc.add_paragraph('بهار ۱۴۰۴')
term_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
term_p.runs[0].font.size = Pt(14)

doc.add_page_break()

# Table of Contents
toc = doc.add_heading('فهرست مطالب', 1)
toc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
toc_list = [
    "۱. چکیده",
    "۲. معماری کلی سیستم",
    "۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)",
    "   ۳.۱. ایجنت‌های هوشمند (Agno Agents)",
    "   ۳.۲. ابزارهای تحلیلی (ToolKits)",
    "   ۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)",
    "۴. معماری فرانت‌اند",
    "   ۴.۱. پشته فناوری (Tech Stack)",
    "   ۴.۲. کامپوننت‌های مهم و کلاس‌ها",
    "۵. نتیجه‌گیری"
]
for item in toc_list:
    p = doc.add_paragraph(item)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_page_break()

def add_heading_rtl(text, level):
    h = doc.add_heading(text, level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return h

def add_paragraph_rtl(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p

# 1. Abstract
add_heading_rtl('۱. چکیده', 1)
add_paragraph_rtl(
    'این گزارش به بررسی معماری، طراحی و پیاده‌سازی پروژه دستیار هوشمند بازار کریپتو (PC_BOT_django) می‌پردازد. '
    'تمرکز اصلی این پروژه بر روی ایجاد یک معماری مبتنی بر ایجنت (Agentic Architecture) با استفاده از فریم‌ورک Agno است '
    'تا بتواند به صورت هوشمندانه به سوالات عمومی، تحلیل‌های تکنیکال و داده‌های درون زنجیره‌ای (On-Chain) پاسخ دهد. '
    'همچنین در بخش فرانت‌اند، از فریم‌ورک Astro برای پیاده‌سازی یک رابط کاربری مدرن، تاریک و راست‌چین (Persian RTL) استفاده شده است.'
)

# 2. System Architecture Overview
add_heading_rtl('۲. معماری کلی سیستم', 1)
add_paragraph_rtl(
    'سیستم از دو بخش اصلی تشکیل شده است: یک بک‌اند قدرتمند مبتنی بر پایتون و جنگو/Agno، و یک فرانت‌اند سریع و مدرن بر پایه Astro. '
    'ارتباط بین این دو بخش از طریق APIهای RESTful (به طور خاص مسیر /api/chat) صورت می‌گیرد. درخواست‌های کاربر دریافت شده، '
    'داده‌های متصل شده (در صورت وجود) پردازش شده و سپس به ایجنت‌های مربوطه جهت تولید پاسخ تحویل داده می‌شوند.'
)

# 3. Backend & AI Architecture
add_heading_rtl('۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)', 1)
add_paragraph_rtl(
    'هسته هوش مصنوعی این پروژه به جای استفاده از روش‌های ساده و یکپارچه، از رویکرد ایجنت‌محور با بهره‌گیری از فریم‌ورک Agno نسخه 2.1.4 استفاده می‌کند. '
    'این روش باعث می‌شود تا وظایف پیچیده تحلیلی شکسته شده و به ایجنت‌های تخصصی واگذار شوند.'
)

add_heading_rtl('۳.۱. ایجنت‌های هوشمند (Agno Agents)', 2)
add_paragraph_rtl(
    'در این سیستم چندین ایجنت تخصصی تعریف شده است:\n'
    '- Technical Agent: متخصص تحلیل نمودارها، کندل‌ها و اندیکاتورها (مانند RSI, MACD, Aroon). این ایجنت با دریافت داده‌های سری زمانی، وضعیت بازار را تحلیل کرده و سیگنال ارائه می‌دهد.\n'
    '- Fundamental Agent: متخصص بررسی داده‌های آن‌چین (On-Chain) مانند جریان پول هوشمند (Smart Money Flow) و رفتار ماینرها.\n'
    '- Code Agent: ایجنت برنامه‌نویس برای تولید اسکریپت‌های تحلیلی پایتون (مثلاً با Pandas) جهت محاسبه اندیکاتورها.\n'
    '- Reactive Agent (Web): ایجنت عمومی با قابلیت جستجو در وب (از طریق DuckDuckGo) برای پاسخ به سوالات عمومی و اخبار کریپتو.'
)

add_heading_rtl('۳.۲. ابزارهای تحلیلی (ToolKits)', 2)
add_paragraph_rtl(
    'هر ایجنت برای انجام وظایف خود به مجموعه‌ای از ابزارها دسترسی دارد. این ابزارها در قالب ToolKit توسعه داده شده‌اند:\n'
    '- TAToolKit: شامل توابعی مانند get_chart_summary، get_recent_candles و get_indicator_data که داده‌های خام را برای ایجنت تکنیکال پردازش می‌کنند.\n'
    '- FundamentalToolKit: شامل توابعی مانند get_fundamental_summary و get_metric_data که برای تحلیل جریان‌های ورودی و خروجی صرافی‌ها و رفتار ماینرها استفاده می‌شوند.\n'
    '- WebTools: دسترسی به موتور جستجوی وب برای استخراج اطلاعات بلادرنگ.'
)

add_heading_rtl('۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)', 2)
add_paragraph_rtl(
    'کامپوننت orchestrator.py وظیفه مسیریابی (Routing) درخواست‌ها را بر عهده دارد. تابع run_agno_chat با بررسی payload ورودی (مثلاً وجود کلید analysis یا analysis_fundamental) '
    'درخواست کاربر را به ایجنت مناسب هدایت می‌کند. همچنین این بخش وظیفه محاسبه دقیق توکن‌های مصرفی و هزینه نهایی پردازش را با توجه به مدل انتخاب شده (مانند gpt-4o, gpt-4.5-preview, o1) انجام می‌دهد.'
)

# 4. Frontend Architecture
add_heading_rtl('۴. معماری فرانت‌اند', 1)

add_heading_rtl('۴.۱. پشته فناوری (Tech Stack)', 2)
add_paragraph_rtl(
    'فرانت‌اند این سیستم به طور کامل با فریم‌ورک Astro و با استفاده از آداپتور @astrojs/node برای حالت Server-Side Rendering (SSR) پیاده‌سازی شده است. '
    'این انتخاب به دلیل سرعت بالای بارگذاری و امکان اجرای کدهای سمت سرور (جهت برقراری ارتباط امن با بک‌اند) صورت گرفته است. استایل‌دهی نیز با استفاده از '
    'CSS خام مدرن با تم تاریک (Dark Mode) و کلاس‌های شیشه‌ای (Glassmorphism) انجام شده است.'
)

add_heading_rtl('۴.۲. کامپوننت‌های مهم و کلاس‌ها', 2)
add_paragraph_rtl(
    'مهم‌ترین بخش‌های پیاده‌سازی شده در فرانت‌اند عبارتند از:\n'
    '- ChatInterface.astro: اصلی‌ترین کامپوننت که شامل فرم ورودی، انتخابگر مدل (Model Selector)، تب‌های انتخاب حالت (عمومی، تکنیکال، بنیادین، کدنویسی) و نمایش پیام‌ها است.\n'
    '- ChatMessage.astro: کامپوننت نمایش حباب‌های چت کاربر و ایجنت. این کامپوننت کلاس‌های راست‌چین (rtl-text) را برای متون فارسی اعمال کرده و قیمت هر درخواست را نیز نشان می‌دهد.\n'
    '- Sidebar.astro: سایدبار پروژه که شامل دکمه ایجاد چت جدید و کارت‌های سناریوهای آماده (تحلیل BTC 1m، جریان پول هوشمند، رفتار ماینرها) با استایل‌های کنتراست بالا (High Contrast) می‌باشد.\n'
    '- api/chat.ts: اندپوینت سمت سرور در Astro که وظیفه برقراری ارتباط (Proxy) با بک‌اند هوش مصنوعی (Agno Orchestrator) را بر عهده دارد.'
)

# 5. Conclusion
add_heading_rtl('۵. نتیجه‌گیری', 1)
add_paragraph_rtl(
    'با مهاجرت از یک سیستم یکپارچه ساده به معماری ایجنت‌محور بر پایه Agno در بک‌اند و استفاده از رابط کاربری سریع مبتنی بر Astro در فرانت‌اند، '
    'اکنون سیستم توانایی پردازش هوشمندانه‌تر درخواست‌ها، انتخاب ابزارهای مناسب برای هر نوع پرسش و ارائه تجربه کاربری بومی (فارسی و دارک مود) '
    'با کنتراست بالا را دارا می‌باشد.'
)

doc.save('PC_BOT_django_Report.docx')
print("Report generated successfully as PC_BOT_django_Report.docx")
