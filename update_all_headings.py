import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define color palette
PINK_COLOR = RGBColor(233, 30, 99)       # Vibrant Pink (#E91E63) for Heading 1
PINK_HEX = "E91E63"

HEADING2_COLOR = RGBColor(194, 24, 91)   # Deep Rose/Magenta (#C2185B) for Heading 2
HEADING2_HEX = "C2185B"

HEADING3_COLOR = RGBColor(106, 27, 154)  # Dark Purple (#6A1B9A) for Heading 3
HEADING3_HEX = "6A1B9A"

BODY_COLOR = RGBColor(33, 37, 41)        # Charcoal (#212529) for Normal text
BODY_HEX = "212529"

FONT_NAME = "Vazirmatn"

def set_paragraph_formatting(p, level=None):
    pPr = p._element.get_or_add_pPr()
    
    # Enable RTL (Right-to-Left)
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    if level is not None:
        # Set Outline Level for MS Word Navigation Pane & TOC
        outlineLvl = OxmlElement('w:outlineLvl')
        outlineLvl.set(qn('w:val'), str(level - 1))
        pPr.append(outlineLvl)
        
        # Keep heading with next paragraph
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

def format_run_font(run, font_name=FONT_NAME, size_pt=None, bold=None, color_rgb=None, color_hex=None):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb

    rPr = run._element.get_or_add_rPr()
    
    # Enable RTL for run
    bidi = OxmlElement('w:rtl')
    bidi.set(qn('w:val'), '1')
    rPr.append(bidi)
    
    # Set rFonts XML attributes (ascii, hAnsi, cs, eastAsia)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

    # Set explicit XML color hex
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_formatting(h, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        format_run_font(run, font_name=FONT_NAME, size_pt=18, bold=True, color_rgb=PINK_COLOR, color_hex=PINK_HEX)
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_formatting(h, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        format_run_font(run, font_name=FONT_NAME, size_pt=15, bold=True, color_rgb=HEADING2_COLOR, color_hex=HEADING2_HEX)
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(text, level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_formatting(h, level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        format_run_font(run, font_name=FONT_NAME, size_pt=13, bold=True, color_rgb=HEADING3_COLOR, color_hex=HEADING3_HEX)
    return h

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_formatting(p)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    format_run_font(r, font_name=FONT_NAME, size_pt=11, color_rgb=BODY_COLOR, color_hex=BODY_HEX)
    return p

def create_formatted_report():
    doc = docx.Document()

    # Style defaults
    styles = doc.styles
    
    # Configure Normal Style
    normal_style = styles['Normal']
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_COLOR

    # Configure Heading 1 Style -> PINK
    h1_style = styles['Heading 1']
    h1_style.font.name = FONT_NAME
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = PINK_COLOR

    # Configure Heading 2 Style
    h2_style = styles['Heading 2']
    h2_style.font.name = FONT_NAME
    h2_style.font.size = Pt(15)
    h2_style.font.bold = True
    h2_style.font.color.rgb = HEADING2_COLOR

    # Configure Heading 3 Style
    h3_style = styles['Heading 3']
    h3_style.font.name = FONT_NAME
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.color.rgb = HEADING3_COLOR

    # Configure Title Style
    title_style = styles['Title']
    title_style.font.name = FONT_NAME
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(15, 23, 42)

    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_formatting(title_p)
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(24)
    r_title = title_p.add_run('گزارش معماری و پیاده‌سازی پروژه دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی')
    format_run_font(r_title, font_name=FONT_NAME, size_pt=22, bold=True, color_rgb=RGBColor(15, 23, 42))

    doc.add_paragraph('\n\n')

    prof_p = doc.add_paragraph()
    prof_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_formatting(prof_p)
    r_prof = prof_p.add_run('درس هوش مصنوعی - استاد تاج بخش')
    format_run_font(r_prof, font_name=FONT_NAME, size_pt=16, bold=True, color_rgb=RGBColor(51, 65, 85))

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_formatting(name_p)
    r_name = name_p.add_run('کیانا نصیری')
    format_run_font(r_name, font_name=FONT_NAME, size_pt=14, bold=True, color_rgb=RGBColor(51, 65, 85))

    term_p = doc.add_paragraph()
    term_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_formatting(term_p)
    r_term = term_p.add_run('تابستان ۱۴۰۵')
    format_run_font(r_term, font_name=FONT_NAME, size_pt=13, color_rgb=RGBColor(100, 116, 139))

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading_1(doc, 'فهرست مطالب')
    
    toc_structure = [
        ("۱. چکیده", 1),
        ("۲. معماری کلی سیستم", 1),
        ("۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)", 1),
        ("۳.۱. ایجنت‌های هوشمند (Agno Agents)", 2),
        ("۳.۲. ابزارهای تحلیلی (ToolKits)", 2),
        ("۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)", 2),
        ("۳.۴. پایگاه داده و ذخیره‌سازی داده‌ها (PostgreSQL & Redis)", 2),
        ("۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها (Live Data Auto-Fetching)", 2),
        ("۴. معماری فرانت‌اند", 1),
        ("۴.۱. پشته فناوری (Tech Stack)", 2),
        ("۴.۲. کامپوننت‌های مهم و کلاس‌ها", 2),
        ("۵. تست و ارزیابی سیستم (Testing & Quality Assurance)", 1),
        ("۵.۱. تست‌های واحد و یکپارچه‌سازی بک‌اند (Pytest Suite)", 2),
        ("۵.۲. تست محاسبات تریدینگ و لوریج", 2),
        ("۵.۳. تست آداپتور فرانت‌اند و API", 2),
        ("۶. نتیجه‌گیری", 1)
    ]
    
    for title, level in toc_structure:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_formatting(p)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        indent = "      " if level == 3 else ("   " if level == 2 else "")
        r = p.add_run(indent + title)
        
        color = PINK_COLOR if level == 1 else (HEADING2_COLOR if level == 2 else HEADING3_COLOR)
        format_run_font(
            r,
            font_name=FONT_NAME,
            size_pt=11.5 if level == 1 else (10.5 if level == 2 else 10.0),
            bold=(level == 1 or level == 2),
            color_rgb=color
        )

    doc.add_page_break()

    # ==================== SECTION 1 ====================
    add_heading_1(doc, '۱. چکیده')
    add_body_paragraph(
        doc,
        "این گزارش به بررسی معماری، طراحی و پیاده‌سازی پروژه دستیار هوشمند بازار کریپتو (دستیار هوش مصنوعی تریدینگ بر اساس مدل‌های زبانی) می‌پردازد. "
        "این پروژه پیش‌تر برای پلتفرم parsiancrpto.com طراحی و توسعه داده شده بود و تمرکز اصلی آن بر روی ایجاد یک معماری مبتنی بر ایجنت (Agentic Architecture) "
        "با بهره‌گیری از فریم‌ورک Agno، پایگاه داده PostgreSQL و Redis، و سیستم دریافت خودکار و زنده داده‌های بازار (Live Market & On-Chain Auto-Fetching) است "
        "تا بتواند به صورت هوشمندانه، لحظه‌ای و پایدار به سوالات عمومی، تحلیل‌های تکنیکال و داده‌های درون زنجیره‌ای (On-Chain) پاسخ دهد. "
        "همچنین در بخش فرانت‌اند، از فریم‌ورک Astro برای پیاده‌سازی یک رابط کاربری مدرن، تاریک و راست‌چین استفاده شده است."
    )

    # ==================== SECTION 2 ====================
    add_heading_1(doc, '۲. معماری کلی سیستم')
    add_body_paragraph(
        doc,
        "سیستم از دو بخش اصلی تشکیل شده است: یک بک‌اند قدرتمند مبتنی بر پایتون، جنگو و فریم‌ورک Agno متصل به پایگاه داده PostgreSQL و Redis، و یک فرانت‌اند سریع و مدرن بر پایه Astro. "
        "این سیستم در ابتدا جهت ارائه خدمات تحلیلی هوشمند به کاربران سامانه parsiancrpto.com طراحی و آماده‌سازی گردید. "
        "ارتباط بین این دو بخش از طریق APIهای RESTful (به طور خاص مسیر /api/chat) صورت می‌گیرد. درخواست‌های کاربر دریافت شده، داده‌های متصل شده (در صورت وجود) پردازش شده و پس از ذخیره‌سازی و مدیریت در پایگاه داده، به ایجنت‌های مربوطه جهت تولید پاسخ تحویل داده می‌شوند."
    )

    # ==================== SECTION 3 ====================
    add_heading_1(doc, '۳. معماری بک‌اند و هوش مصنوعی (Agno Framework)')
    add_body_paragraph(
        doc,
        "هسته هوش مصنوعی این پروژه به جای استفاده از روش‌های ساده و یکپارچه، از رویکرد ایجنت‌محور با بهره‌گیری از فریم‌ورک Agno نسخه 2.1.4 استفاده می‌کند. "
        "این روش باعث می‌شود تا وظایف پیچیده تحلیلی شکسته شده و به ایجنت‌های تخصصی واگذار شوند."
    )

    # --- Section 3.1 ---
    add_heading_2(doc, '۳.۱. ایجنت‌های هوشمند (Agno Agents)')
    add_body_paragraph(doc, "در این سیستم چندین ایجنت تخصصی برای انجام وظایف تحلیلی و پردازشی تعریف شده است:")

    add_heading_3(doc, '۳.۱.۱. ایجنت تحلیل تکنیکال (Technical Agent)')
    add_body_paragraph(
        doc,
        "متخصص تحلیل نمودارها، کندل‌های OHLCV و اندیکاتورهای فنی (مانند RSI, MACD, Aroon, CCI). این ایجنت با دریافت داده‌های سری زمانی، وضعیت بازار را تحلیل کرده و سیگنال‌های دقیق معاملاتی صادر می‌نماید."
    )

    add_heading_3(doc, '۳.۱.۲. ایجنت بنیادی و آن‌چین (Fundamental Agent)')
    add_body_paragraph(
        doc,
        "متخصص بررسی داده‌های درون زنجیره‌ای (On-Chain)، شاخص‌های جریان پول هوشمند (Smart Money Flow)، خالص جریان صرافی‌ها (Exchange Netflow) و تحلیل رفتار ماینرها."
    )

    add_heading_3(doc, '۳.۱.۳. ایجنت برنامه‌نویسی و کد (Code Agent)')
    add_body_paragraph(
        doc,
        "ایجنت برنامه‌نویس جهت تولید اسکریپت‌های تحلیلی پایتون (با استفاده از کتابخانه Pandas) برای محاسبه فرمول‌های پیچیده اندیکاتورها و پردازش داده‌ها."
    )

    add_heading_3(doc, '۳.۱.۴. ایجنت عمومی و وب (Reactive Agent)')
    add_body_paragraph(
        doc,
        "ایجنت عمومی با قابلیت جستجوی مستقیم در وب (از طریق موتور جستجوی DuckDuckGo) برای پاسخ به سوالات عمومی، احوال‌پرسی‌ها و استخراج اخبار کریپتو."
    )

    # --- Section 3.2 ---
    add_heading_2(doc, '۳.۲. ابزارهای تحلیلی (ToolKits)')
    add_body_paragraph(doc, "هر ایجنت برای انجام وظایف خود به مجموعه‌ای از ابزارهای تخصصی (ToolKits) دسترسی دارد:")

    add_heading_3(doc, '۳.۲.۱. ابزار تحلیل تکنیکال (TAToolKit)')
    add_body_paragraph(
        doc,
        "شامل توابعی نظیر get_chart_summary، get_recent_candles و get_indicator_data که داده‌های خام کندل‌ها و اندیکاتورها را جهت استفاده ایجنت تکنیکال پردازش می‌کنند."
    )

    add_heading_3(doc, '۳.۲.۲. ابزار داده‌های بنیادی (FundamentalToolKit)')
    add_body_paragraph(
        doc,
        "شامل توابعی مانند get_fundamental_summary و get_metric_data که جهت تحلیل جریان ورودی/خروجی صرافی‌ها و رفتار ماینرها به کار می‌روند."
    )

    add_heading_3(doc, '۳.۲.۳. ابزارهای جستجوی وب (WebTools)')
    add_body_paragraph(
        doc,
        "دسترس‌پذیری به موتور جستجوی وب جهت استخراج اطلاعات بلادرنگ و اخبار روز کریپتو."
    )

    # --- Section 3.3 ---
    add_heading_2(doc, '۳.۳. هماهنگ‌کننده و مسیریابی (Orchestrator)')
    add_body_paragraph(
        doc,
        "کامپوننت orchestrator.py وظیفه مسیریابی (Routing) هوشمند درخواست‌ها را بر عهده دارد. تابع run_agno_chat با بررسی دقیق متن پرسش و دیتاسِت‌های ورودی، درخواست کاربر را به ایجنت تخصصی مربوطه (Technical, Fundamental, Code یا Reactive) هدایت می‌کند. "
        "در این لایه، گام‌های روند تفکر ایجنت (Thinking Process) به صورت شفاف استخراج شده و محاسبه دقیق توکن‌های مصرفی و قیمت نهایی پردازش به دلار (مانند $0.000080 برای gpt-4o-mini) بر اساس فرمول‌های قیمت‌گذاری مدل انجام می‌پذیرد."
    )

    # --- Section 3.4 ---
    add_heading_2(doc, '۳.۴. پایگاه داده و ذخیره‌سازی داده‌ها (PostgreSQL & Redis)')
    add_body_paragraph(
        doc,
        "جهت مدیریت نشست‌های کاری (Conversations)، ذخیره‌سازی پیام‌ها، نگه‌داری متاداده‌های تحلیلی و ثبت بازخوردهای کاربران، بک‌اند سیستم از پایگاه داده رابطه‌ای PostgreSQL به همراه ORM قدرتمند جنگو استفاده می‌کند. مهم‌ترین مدل‌های ذخیره‌سازی طراحی‌شده عبارتند از:"
    )

    add_heading_3(doc, '۳.۴.۱. مدل نشست‌ها (Conversation)')
    add_body_paragraph(
        doc,
        "نگه‌داری شناسه یکتای نشست (UUID)، شناسه کاربر (user_id)، وضعیت آرشیو (is_archive) و زمان‌بندی ایجاد و به‌روزرسانی."
    )

    add_heading_3(doc, '۳.۴.۲. مدل پیام‌ها (Message)')
    add_body_paragraph(
        doc,
        "مدل اصلی ذخیره‌سازی پیام‌ها شامل متن پرسش کاربر (question)، پاسخ ایجنت (answer)، قیمت دلار (price)، مدل انتخاب‌شده (gpt-4o, o3-mini)، وضعیت اجرا (status)، سطح استدلال (reasoning)، مدت زمان پاسخ‌دهی و فیلدهای ساختاریافته JSON (تحت عنوان analysis و analysis_fundamental)."
    )

    add_heading_3(doc, '۳.۴.۳. مدل بازخورد کاربران (Feedback)')
    add_body_paragraph(
        doc,
        "ثبت بازخوردهای کیفی کاربران (لایک / دیس‌لایک) به همراه شناسه کاربر و کلید خارجی متصل به پیام جهت سنجش دقیق کیفیت ایجنت‌ها."
    )

    add_heading_3(doc, '۳.۴.۴. مدل گزارش خطا (ProblemReport)')
    add_body_paragraph(
        doc,
        "سیستم ثبت و پیگیری گزارش خطاهای کاربران به همراه وضعیت بررسی (Open/Close) و پاسخ مدیر سیستم (close_admin_id)."
    )

    add_heading_3(doc, '۳.۴.۵. لایه حافظه پنهان و صف وظایف (Redis)')
    add_body_paragraph(
        doc,
        "پایگاه داده کلید-مقدار Redis به عنوان Broker پیام‌رسان برای مدیریت وظایف پیش‌زمینه (Asynchronous Tasks) و لایه لایه‌بندی حافظه پنهان (Caching) استفاده گردیده است."
    )

    # --- Section 3.5 ---
    add_heading_2(doc, '۳.۵. مکانیزم دریافت خودکار و زنده داده‌ها (Live Data Auto-Fetching)')
    add_body_paragraph(
        doc,
        "یکی از کلیدی‌ترین قابلیت‌های سیستم، عدم اتکا به داده‌های قدیمی و توانایی دریافت خودکار و لحظه‌ای داده‌های بازار کریپتو (Live Market Data) می‌باشد. گام‌های این مکانیزم به شرح زیر است:"
    )

    add_heading_3(doc, '۳.۵.۱. شناسایی هوشمند حالت تحلیل (Query Intent Detection)')
    add_body_paragraph(
        doc,
        "هنگام طرح پرسش‌های تحلیلی، سرویس MarketDataService کلیدواژه‌های تخصصی تکنیکال (مانند RSI, MACD, کندل، قیمت) و آن‌چین (مانند پول هوشمند، رفتار ماینرها) را بررسی می‌کند."
    )

    add_heading_3(doc, '۳.۵.۲. فراخوانی APIهای لحظه‌ای (Live Market Data API)')
    add_body_paragraph(
        doc,
        "در صورت نیاز به داده، داده‌های لحظه‌ای قیمت و کندل‌های صرافی Binance (پروتوکل API klines) و شاخص‌های آن‌چین CoinGecko استخراج می‌شوند."
    )

    add_heading_3(doc, '۳.۵.۳. بهینه‌سازی عدم فراخوانی در گفتگوهای عمومی (General Chat Optimization)')
    add_body_paragraph(
        doc,
        "در نسخه قبلی، عدم کشف کلیدواژه باعث فراخوانی پیش‌فرض داده‌های بیت‌کوین برای پرسش‌های ساده (مانند 'hi' یا 'سلام') می‌شد. با بهینه‌سازی صورت‌گرفته در تابع auto_enrich_payloads و افزودن تست واحد test_auto_enrich_payloads_greeting، برای پیام‌های عمومی هیچ فراخوانی API خارجی انجام نگرفته و پرسش مستقیماً توسط Reactive Agent پاسخ داده می‌شود."
    )

    add_heading_3(doc, '۳.۵.۴. تزریق به ابزارهای ایجنت (ToolKit Injection)')
    add_body_paragraph(
        doc,
        "داده‌های زنده به صورت ساختاریافته (JSON) به ابزارهای تحلیلی ایجنت تزریق می‌شوند تا ایجنت Agno سطوح حمایت/مقاومت، روند بازار و سیگنال‌های معاملاتی (EP, TP, SL, R:R) را استخراج کند."
    )

    # ==================== SECTION 4 ====================
    add_heading_1(doc, '۴. معماری فرانت‌اند')

    add_heading_2(doc, '۴.۱. پشته فناوری (Tech Stack)')
    add_body_paragraph(
        doc,
        "فرانت‌اند این سیستم به طور کامل با فریم‌ورک Astro و با استفاده از آداپتور @astrojs/node برای حالت Server-Side Rendering (SSR) پیاده‌سازی شده است. "
        "این انتخاب به دلیل سرعت بالای بارگذاری و امکان اجرای کدهای سمت سرور (جهت برقراری ارتباط امن با بک‌اند) صورت گرفته است. استایل‌دهی نیز با استفاده از "
        "CSS خام مدرن با تم تاریک (Dark Mode) و کلاس‌های شیشه‌ای (Glassmorphism) انجام شده است."
    )

    add_heading_2(doc, '۴.۲. کامپوننت‌های مهم و کلاس‌ها')
    add_body_paragraph(doc, "مهم‌ترین بخش‌های پیاده‌سازی شده در فرانت‌اند عبارتند از:")

    add_heading_3(doc, '۴.۲.۱. کامپوننت رابط کاربری چت (ChatInterface.astro)')
    add_body_paragraph(
        doc,
        "اصلی‌ترین کامپوننت که شامل فرم ورودی، انتخابگر مدل (Model Selector)، تب‌های انتخاب حالت (عمومی، تکنیکال، بنیادین، کدنویسی) و نمایش پیام‌ها است."
    )

    add_heading_3(doc, '۴.۲.۲. کامپوننت نمایش پیام (ChatMessage.astro)')
    add_body_paragraph(
        doc,
        "کامپوننت نمایش حباب‌های چت کاربر و ایجنت. این کامپوننت کلاس‌های راست‌چین (rtl-text) را برای متون فارسی اعمال کرده و قیمت هر درخواست را نیز نشان می‌دهد."
    )

    add_heading_3(doc, '۴.۲.۳. کامپوننت نوار کناری (Sidebar.astro)')
    add_body_paragraph(
        doc,
        "سایدبار پروژه که شامل دکمه ایجاد چت جدید و کارت‌های سناریوهای آماده (تحلیل BTC 1m، جریان پول هوشمند، رفتار ماینرها) با استایل‌های کنتراست بالا (High Contrast) می‌باشد."
    )

    add_heading_3(doc, '۴.۲.۴. اندپوینت API سمت سرور (api/chat.ts)')
    add_body_paragraph(
        doc,
        "اندپوینت سمت سرور در Astro که وظیفه برقراری ارتباط (Proxy) با بک‌اند هوش مصنوعی (Agno Orchestrator) را بر عهده دارد."
    )

    # ==================== SECTION 5 ====================
    add_heading_1(doc, '۵. تست و ارزیابی سیستم (Testing & Quality Assurance)')

    add_heading_2(doc, '۵.۱. تست‌های واحد و یکپارچه‌سازی بک‌اند (Pytest Suite)')
    add_body_paragraph(
        doc,
        "جهت حصول اطمینان از صحت عملکرد ایجنت‌ها، ابزارهای تحلیلی و کامپوننت هماهنگ‌کننده، یک مجموعه جامع از تست‌های واحد و یکپارچه‌سازی با فریم‌ورک Pytest پیاده‌سازی شده است (شامل ۱۶ تست با موفقیت ۱۰۰٪). مهم‌ترین بخش‌های تست شده عبارتند از:"
    )

    add_heading_3(doc, '۵.۱.۱. تست‌های ورودی چت و ایجنت (test_chat_bot.py)')
    add_body_paragraph(
        doc,
        "ارزیابی تابع اصلی ورودی چت، مدیریت ورودی‌های خالی و شبیه‌سازی (Mocking) پاسخ‌های ایجنت Agno شامل محاسبه توکن‌ها و هزینه درخواست."
    )

    add_heading_3(doc, '۵.۱.۲. تست ابزارهای بنیادی (test_fundamental_tools.py)')
    add_body_paragraph(
        doc,
        "تست ابزارهای آن‌چین در FundamentalToolKit مانند خلاصه متریال‌های بنیادی، متدهای لیست‌گیری و استخراج داده‌های جریان پول هوشمند و خالص جریان صرافی‌ها."
    )

    add_heading_3(doc, '۵.۱.۳. تست هماهنگ‌کننده و قیمت‌گذاری (test_orchestrator.py)')
    add_body_paragraph(
        doc,
        "تست توابع هماهنگ‌کننده شامل تخمین توکن‌ها، قالب‌بندی بافت تاریخچه مکالمه، محاسبه قیمت دقیق مدل‌های مختلف (gpt-4o-mini, o3-mini) و اجرای حالت جایگزین (Fallback)."
    )

    add_heading_3(doc, '۵.۱.۴. تست ابزارهای تکنیکال (test_ta_tools.py)')
    add_body_paragraph(
        doc,
        "سنجش ابزارهای تحلیل تکنیکال TAToolKit شامل خلاصه نمودار، دریافت کندل‌های OHLCV، محاسبه آمار قیمت و پردازش خروجی اندیکاتورهای RSI, Aroon و MACD."
    )

    add_heading_3(doc, '۵.۱.۵. تست عدم فراخوانی در احوال‌پرسی (test_auto_enrich_payloads_greeting)')
    add_body_paragraph(
        doc,
        "تست تایید عدم فراخوانی APIهای خارجی Binance و CoinGecko هنگام ارسال پیام‌های احوال‌پرسی ساده (مانند 'hi')."
    )

    add_heading_2(doc, '۵.۲. تست محاسبات تریدینگ و لوریج')
    add_body_paragraph(
        doc,
        "یکی از حساس‌ترین بخش‌های پروژه، اعتبارسنجی محاسبات تریدینگ است. در این راستا تست‌های دقیق برای توابع ریاضی زیر اجرا گردید:"
    )

    add_heading_3(doc, '۵.۲.۱. محاسبه ضریب لوریج بر اساس درصد حد زیان')
    add_body_paragraph(doc, "اعتبارسنجی فرمول LV = 90 / %SL و گرد کردن آن به سمت پایین.")

    add_heading_3(doc, '۵.۲.۲. محاسبه نسبت ریسک به ریوارد (Risk-Reward Ratio)')
    add_body_paragraph(doc, "ارزیابی نسبت سود به زیان برای پوزیشن‌های خرید (Long) و فروش (Short) بر اساس قیمت ورود، حد زیان و اهداف سود (TP1..N).")

    add_heading_3(doc, '۵.۲.۳. اعتبارسنجی پارامترهای معاملاتی (validate_trading_calculations)')
    add_body_paragraph(doc, "کنترل حد زیان متناسب با تایم‌فریم (مانند حداقل ۲٪ برای تایم‌فریم ۱۵ دقیقه) و صدور هشدار در صورت ریسک نامناسب.")

    add_heading_2(doc, '۵.۳. تست آداپتور فرانت‌اند و API')
    add_body_paragraph(
        doc,
        "در بخش فرانت‌اند Astro، اندپوینت api/chat.ts از نظر سرعت پاسخ‌دهی سمت سرور (SSR)، لایه‌بندی حباب‌های پیام، نمایش کارت‌های سناریو پیش‌فرض و مدیریت خطاهای شبکه مورد ارزیابی قرار گرفت تا تجربه کاربری روان و بلادرنگ برای کاربران تضمین گردد."
    )

    # ==================== SECTION 6 ====================
    add_heading_1(doc, '۶. نتیجه‌گیری')
    add_body_paragraph(
        doc,
        "با مهاجرت از یک سیستم یکپارچه ساده به معماری ایجنت‌محور بر پایه Agno در بک‌اند، بهره‌گیری از پایگاه داده PostgreSQL و Redis جهت مدیریت مطمئن نشست‌ها و کارهای پیش‌زمینه، و استفاده از رابط کاربری سریع مبتنی بر Astro در فرانت‌اند، اکنون سیستم توانایی پردازش هوشمندانه‌تر درخواست‌ها، انتخاب ابزارهای مناسب برای هر نوع پرسش و ارائه تجربه کاربری بومی (فارسی و دارک مود) با کنتراست بالا را دارا می‌باشد."
    )

    doc.save('Bachelor_proj_report.docx')
    print("Report formatted successfully with ALL sub-sections and sub-items as proper Heading formats!")

if __name__ == '__main__':
    create_formatted_report()
